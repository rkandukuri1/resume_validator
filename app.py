from fastapi import FastAPI, UploadFile, File, Form
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import FileResponse
from openai import OpenAI
from dotenv import load_dotenv
from pydantic import BaseModel, Field
from typing import List
import docx2txt
from docx import Document
import tempfile
import shutil
import os

# -----------------------------
# Load API Key
# -----------------------------
load_dotenv()
client = OpenAI(api_key=os.getenv("OPENAI_API_KEY_FREE"))

# -----------------------------
# Create App
# -----------------------------
app = FastAPI(title="Resume Validator API")

app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

# -----------------------------
# Pydantic Models
# -----------------------------
class ResumeValidator(BaseModel):
    match_score: int = Field(ge=0, le=100)
    matched_skills: List[str]
    missing_skills: List[str]
    experience_match: str
    strengths: List[str]
    suggestions: List[str]
    final_resolution: str
    hiring_status: str


class ImprovedResume(BaseModel):
    improved_resume: str


# -----------------------------
# Resume Validation
# -----------------------------
def process_resume(resume_docs, req_docs):

    prompt = f"""
You are an expert resume validator.

Validate the resume against the requirement.

- Penalize missing required skills
- Same input must produce nearly same score every time
- Be strict and objective

Provide:
- match_score
- matched_skills
- missing_skills
- strengths
- suggestions
- experience_match
- final_resolution

Hiring status:
Consider = 75+
WaitList = 50-74
Reject = below 50

Match score must be between 0 and 100.
At the end provide hiring status like:
Status : Consider / WaitList / Reject

RESUME:
{resume_docs}

REQUIREMENT:
{req_docs}
"""

    response = client.beta.chat.completions.parse(
        messages=[{"role": "user", "content": prompt}],
        model="gpt-4.1-nano",
        temperature=0,
        top_p=0.1,
        response_format=ResumeValidator
    )

    return response.choices[0].message.parsed


# -----------------------------
# Resume Improver
# -----------------------------
def generate_improved_resume(resume_text, req_text, missing_skills):

    prompt = f"""
You are an expert resume writer.

Improve the resume so it better matches the job requirement.

Rules:
- Keep original structure meaning what ever are Bold, Italics and Underlines in main resume
must be accounted and implemented in new resume as is.
- Add missing skills naturally
- Add responsibilities where needed
- Keep experience realistic
- Make it ATS friendly
- Dont insert extra blank lines.
- Keep spacing compact.
- Preserve existing paragraph structure.
- Preserve bullet formatting
- Every responsibility must remain a separate bullet
- Never combine multiple bullet points into one bullet
- Each bullet should contain only one achievement
- Keep formatting compact
 
Missing Skills:
{missing_skills}

REQUIREMENT:
{req_text}

ORIGINAL RESUME:
{resume_text}
"""

    response = client.beta.chat.completions.parse(
        messages=[{"role": "user", "content": prompt}],
        model="gpt-4.1",
        temperature=0,
        top_p=0.1,
        response_format=ImprovedResume
    )

    return response.choices[0].message.parsed


# -----------------------------
# Create Word File
# -----------------------------
def create_resume_file(original_file, new_text):

    original_doc = Document(original_file)
    new_doc = Document()

    new_lines = [line.strip() for line in new_text.split("\n") if line.strip()]

    for i, line in enumerate(new_lines):

        source_para = original_doc.paragraphs[i] if i < len(original_doc.paragraphs) else None
        is_bullet = line.startswith(("•", "-", "*"))

        if is_bullet:
            clean_line = line.lstrip("•-* ").strip()
            new_para = new_doc.add_paragraph(style="List Bullet")
        else:
            clean_line = line
            new_para = new_doc.add_paragraph()

            if source_para:
                new_para.style = source_para.style

        if source_para and source_para.runs:
            src_run = source_para.runs[0]

            run = new_para.add_run(clean_line)
            run.bold = src_run.bold
            run.italic = src_run.italic
            run.underline = src_run.underline
            run.font.name = src_run.font.name
            run.font.size = src_run.font.size
        else:
            new_para.add_run(clean_line)

    file_path = "improved_resume.docx"
    new_doc.save(file_path)

    return file_path


# -----------------------------
# Home
# -----------------------------
@app.get("/")
def home():
    return {"message": "Resume Validator API running"}


# -----------------------------
# Validate Resume
# -----------------------------
@app.post("/validate-resume")
async def validate_resume(
    resume_file: UploadFile = File(...),
    requirement_text: str = Form(...)
):
    with tempfile.NamedTemporaryFile(delete=False, suffix=".docx") as tmp:
        shutil.copyfileobj(resume_file.file, tmp)
        temp_path = tmp.name

    resume_docs = docx2txt.process(temp_path)

    result = process_resume(resume_docs, requirement_text)

    return result.model_dump()


# -----------------------------
# Improve Resume
# -----------------------------
@app.post("/improve-resume")
async def improve_resume(
    resume_file: UploadFile = File(...),
    requirement_text: str = Form(...)
):
    with tempfile.NamedTemporaryFile(delete=False, suffix=".docx") as tmp:
        shutil.copyfileobj(resume_file.file, tmp)
        temp_path = tmp.name

    resume_docs = docx2txt.process(temp_path)

    result = process_resume(resume_docs, requirement_text)

    improved = generate_improved_resume(
        resume_docs,
        requirement_text,
        result.missing_skills
    )

    file_path = create_resume_file(temp_path, improved.improved_resume)

    return FileResponse(
        path=file_path,
        filename="improved_resume.docx",
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )
    st.warning("File Downloaded to default Downloads folder")
    st.warning("There might be some format issues in the Improved file like Bold/Underline/Italics etc. Please check.") 